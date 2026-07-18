"""
مسارات النصوص المفرَّغة — جلب، تعديل، ترجمة، تلخيص، تحديد متحدثين، فصول ذكية
"""
import json
import io
from typing import Optional, List
from fastapi import APIRouter, Depends, HTTPException, BackgroundTasks
from fastapi.responses import StreamingResponse, PlainTextResponse, JSONResponse
from sqlalchemy.orm import Session
from pydantic import BaseModel
from datetime import datetime

from app.database import get_db, Transcript, Video, TranscriptStatus, User
from app.auth import get_current_user, require_auth

router = APIRouter()


# ── Schemas ───────────────────────────────────────────
class SegmentSchema(BaseModel):
    start:   float
    end:     float
    text:    str
    speaker: Optional[str] = None
    words:   Optional[list] = None

class TranscriptResponse(BaseModel):
    id:                str
    video_id:          str
    full_text:         Optional[str]
    segments:          Optional[List[SegmentSchema]]
    status:            str
    language_detected: Optional[str]
    processing_time:   Optional[float]
    error_message:     Optional[str]
    updated_at:        datetime

    class Config:
        from_attributes = True

class EditTranscriptRequest(BaseModel):
    full_text: Optional[str] = None
    segments:  Optional[list] = None

class ChapterSchema(BaseModel):
    start:   float
    end:     float
    title:   str
    summary: str


# ── Helper ────────────────────────────────────────────
def _get_transcript_or_404(video_id, db, current_user=None):
    video = db.query(Video).filter(Video.id == video_id).first()
    if not video:
        raise HTTPException(404, "الفيديو غير موجود")
    if not video.is_public:
        if not current_user or video.owner_id != current_user.id:
            raise HTTPException(403, "غير مصرح")
    transcript = db.query(Transcript).filter(Transcript.video_id == video_id).first()
    if not transcript:
        raise HTTPException(404, "لا يوجد تفريغ")
    return transcript


# ══════════════════════════════════════════════════════
#  GET /api/transcripts/{video_id}
# ══════════════════════════════════════════════════════
@router.get("/{video_id}", response_model=TranscriptResponse)
def get_transcript(
    video_id:     str,
    db:           Session        = Depends(get_db),
    current_user: Optional[User] = Depends(get_current_user),
):
    transcript = _get_transcript_or_404(video_id, db, current_user)
    segments = json.loads(transcript.segments_json) if transcript.segments_json else None
    return TranscriptResponse(
        id=transcript.id, video_id=transcript.video_id,
        full_text=transcript.full_text, segments=segments,
        status=transcript.status, language_detected=transcript.language_detected,
        processing_time=transcript.processing_time,
        error_message=transcript.error_message, updated_at=transcript.updated_at,
    )


# ══════════════════════════════════════════════════════
#  PATCH /api/transcripts/{video_id}
# ══════════════════════════════════════════════════════
@router.patch("/{video_id}", response_model=TranscriptResponse)
def edit_transcript(
    video_id:     str,
    data:         EditTranscriptRequest,
    db:           Session = Depends(get_db),
    current_user: User    = Depends(require_auth),
):
    video = db.query(Video).filter(Video.id == video_id, Video.owner_id == current_user.id).first()
    if not video:
        raise HTTPException(404, "غير موجود أو ليس لديك صلاحية")
    transcript = db.query(Transcript).filter(Transcript.video_id == video_id).first()
    if not transcript:
        raise HTTPException(404, "لا يوجد تفريغ")
    if data.full_text is not None:
        transcript.full_text = data.full_text
    if data.segments is not None:
        transcript.segments_json = json.dumps(data.segments, ensure_ascii=False)
    db.commit(); db.refresh(transcript)
    segments = json.loads(transcript.segments_json) if transcript.segments_json else None
    return TranscriptResponse(
        id=transcript.id, video_id=transcript.video_id,
        full_text=transcript.full_text, segments=segments,
        status=transcript.status, language_detected=transcript.language_detected,
        processing_time=transcript.processing_time,
        error_message=transcript.error_message, updated_at=transcript.updated_at,
    )


# ══════════════════════════════════════════════════════
#  POST /api/transcripts/{video_id}/retry
# ══════════════════════════════════════════════════════
@router.post("/{video_id}/retry")
def retry_transcription(
    video_id:         str,
    background_tasks: BackgroundTasks,
    db:               Session = Depends(get_db),
    current_user:     User    = Depends(require_auth),
):
    video = db.query(Video).filter(Video.id == video_id, Video.owner_id == current_user.id).first()
    if not video:
        raise HTTPException(404, "الفيديو غير موجود")
    transcript = db.query(Transcript).filter(Transcript.video_id == video_id).first()
    if transcript.status == TranscriptStatus.PROCESSING:
        raise HTTPException(400, "التفريغ قيد المعالجة")
    transcript.status = TranscriptStatus.PENDING
    transcript.error_message = None
    db.commit()
    from app.routers.videos import run_transcription_task
    background_tasks.add_task(
        run_transcription_task,
        video_id=video_id,
        file_path=video.file_path,
        language=video.dialect if len(video.dialect) == 2 else "ar",
    )
    return {"message": "تمت جدولة إعادة التفريغ", "status": "pending"}


# ══════════════════════════════════════════════════════
#  POST /api/transcripts/{video_id}/translate  🌍
# ══════════════════════════════════════════════════════
@router.post("/{video_id}/translate")
def translate_transcript(
    video_id:     str,
    db:           Session        = Depends(get_db),
    current_user: Optional[User] = Depends(get_current_user),
):
    """ترجمة النص المفرَّغ من العربية للإنجليزية"""
    transcript = _get_transcript_or_404(video_id, db, current_user)

    if transcript.status != TranscriptStatus.DONE:
        raise HTTPException(400, "التفريغ لم يكتمل بعد")
    if not transcript.full_text:
        raise HTTPException(400, "لا يوجد نص للترجمة")

    try:
        from app.ai_services import translate_to_english
        segments = json.loads(transcript.segments_json) if transcript.segments_json else []
        result   = translate_to_english(transcript.full_text, segments)
        return {
            "full_text_en": result.get("full_text", ""),
            "segments_en":  result.get("segments", []),
            "source_lang":  "ar",
            "target_lang":  "en",
        }
    except ValueError as e:
        raise HTTPException(400, str(e))
    except Exception as e:
        raise HTTPException(500, f"فشلت الترجمة: {str(e)}")


# ══════════════════════════════════════════════════════
#  POST /api/transcripts/{video_id}/summarize  🤖
# ══════════════════════════════════════════════════════
@router.post("/{video_id}/summarize")
def summarize_transcript_route(
    video_id:     str,
    db:           Session        = Depends(get_db),
    current_user: Optional[User] = Depends(get_current_user),
):
    """تلخيص النص بالذكاء الاصطناعي"""
    transcript = _get_transcript_or_404(video_id, db, current_user)

    if transcript.status != TranscriptStatus.DONE:
        raise HTTPException(400, "التفريغ لم يكتمل بعد")
    if not transcript.full_text:
        raise HTTPException(400, "لا يوجد نص للتلخيص")

    try:
        from app.ai_services import summarize_transcript
        result = summarize_transcript(
            transcript.full_text,
            language=transcript.language_detected or "ar"
        )
        transcript.summary = json.dumps(result, ensure_ascii=False)
        db.commit()
        return result
    except ValueError as e:
        raise HTTPException(400, str(e))
    except Exception as e:
        raise HTTPException(500, f"فشل التلخيص: {str(e)}")


# ══════════════════════════════════════════════════════
#  POST /api/transcripts/{video_id}/diarize  👥
# ══════════════════════════════════════════════════════
@router.post("/{video_id}/diarize")
def diarize_transcript(
    video_id:         str,
    background_tasks: BackgroundTasks,
    num_speakers:     Optional[int] = None,
    db:               Session       = Depends(get_db),
    current_user:     User          = Depends(require_auth),
):
    """تحديد المتحدثين في التسجيل"""
    video = db.query(Video).filter(Video.id == video_id, Video.owner_id == current_user.id).first()
    if not video:
        raise HTTPException(404, "الفيديو غير موجود")
    transcript = db.query(Transcript).filter(Transcript.video_id == video_id).first()
    if not transcript or transcript.status != TranscriptStatus.DONE:
        raise HTTPException(400, "التفريغ لم يكتمل بعد")

    try:
        from app.ai_services import diarize_audio, merge_diarization_with_transcript
        segments = json.loads(transcript.segments_json or "[]")

        speakers = diarize_audio(video.file_path, num_speakers)
        merged = merge_diarization_with_transcript(segments, speakers)

        transcript.segments_json = json.dumps(merged, ensure_ascii=False)
        db.commit()

        return {
            "message":         "تم تحديد المتحدثين بنجاح",
            "speakers_found":  len(set(s["speaker"] for s in speakers)),
            "segments":        merged,
        }
    except (ImportError, ValueError) as e:
        raise HTTPException(400, str(e))
    except Exception as e:
        raise HTTPException(500, f"فشل تحديد المتحدثين: {str(e)}")


# ══════════════════════════════════════════════════════
#  POST /api/transcripts/{video_id}/chapters  📑 Feature 2
# ══════════════════════════════════════════════════════
@router.post("/{video_id}/chapters")
def generate_chapters(
    video_id:     str,
    db:           Session        = Depends(get_db),
    current_user: Optional[User] = Depends(get_current_user),
):
    """يُنشئ فصولاً ذكية من النص المفرَّغ باستخدام Claude"""
    transcript = _get_transcript_or_404(video_id, db, current_user)

    if transcript.status != TranscriptStatus.DONE:
        raise HTTPException(400, "التفريغ لم يكتمل بعد")
    if not transcript.full_text:
        raise HTTPException(400, "لا يوجد نص لإنشاء الفصول")

    try:
        import anthropic
        import os

        api_key = os.getenv("ANTHROPIC_API_KEY")
        if not api_key:
            raise HTTPException(500, "ANTHROPIC_API_KEY غير موجود")

        segments = json.loads(transcript.segments_json) if transcript.segments_json else []
        segments_text = "\n".join([
            f"[{s.get('start', 0):.1f}s - {s.get('end', 0):.1f}s]: {s.get('text', '')}"
            for s in segments
        ])

        prompt = f"""أنت مساعد ذكي. لديك نص مفرَّغ من تسجيل صوتي/فيديو.
قم بتقسيمه إلى فصول (chapters) منطقية.

أرجع JSON فقط بدون أي نص إضافي:
{{"chapters": [{{"start": 0.0, "end": 120.5, "title": "عنوان الفصل", "summary": "ملخص الفصل في جملة"}}]}}

- كل فصل يجب أن يكون منطقياً في المحتوى
- العنوان بالعربية
- استخدم الطوابع الزمنية الفعلية من النص
- لا تُخترع توقيتاً غير موجود في النص

النص:
{segments_text[:8000]}"""

        client = anthropic.Anthropic(api_key=api_key)
        message = client.messages.create(
            model="claude-sonnet-4-6",
            max_tokens=2048,
            messages=[{"role": "user", "content": prompt}]
        )

        raw = message.content[0].text.strip()
        raw = raw.replace("```json", "").replace("```", "").strip()
        chapters_data = json.loads(raw)

        # احفظ في قاعدة البيانات
        transcript.chapters_json = json.dumps(chapters_data, ensure_ascii=False)
        db.commit()

        return chapters_data

    except json.JSONDecodeError:
        raise HTTPException(500, "فشل تحليل استجابة الذكاء الاصطناعي")
    except Exception as e:
        raise HTTPException(500, f"فشل إنشاء الفصول: {str(e)}")


# ══════════════════════════════════════════════════════
#  GET /api/transcripts/{video_id}/chapters
# ══════════════════════════════════════════════════════
@router.get("/{video_id}/chapters")
def get_chapters(
    video_id:     str,
    db:           Session        = Depends(get_db),
    current_user: Optional[User] = Depends(get_current_user),
):
    """جلب الفصول المحفوظة"""
    transcript = _get_transcript_or_404(video_id, db, current_user)
    if transcript.chapters_json:
        return json.loads(transcript.chapters_json)
    return {"chapters": []}


# ══════════════════════════════════════════════════════
#  GET /api/transcripts/{video_id}/export  📥
# ══════════════════════════════════════════════════════
@router.get("/{video_id}/export")
def export_transcript(
    video_id:     str,
    fmt:          str           = "txt",
    db:           Session       = Depends(get_db),
    current_user: Optional[User]= Depends(get_current_user),
):
    """تصدير النص: txt, srt, json, docx"""
    transcript = _get_transcript_or_404(video_id, db, current_user)
    if transcript.status != TranscriptStatus.DONE:
        raise HTTPException(400, "التفريغ لم يكتمل بعد")

    video    = db.query(Video).filter(Video.id == video_id).first()
    segments = json.loads(transcript.segments_json or "[]")

    if fmt == "txt":
        return PlainTextResponse(
            content=transcript.full_text or "",
            media_type="text/plain; charset=utf-8",
            headers={"Content-Disposition": f'attachment; filename="transcript.txt"'},
        )

    elif fmt == "srt":
        lines = []
        for i, seg in enumerate(segments, 1):
            lines.append(f"{i}\n{_srt_time(seg['start'])} --> {_srt_time(seg['end'])}\n{seg.get('speaker', '')}{': ' if seg.get('speaker') else ''}{seg['text']}\n")
        return PlainTextResponse(
            content="\n".join(lines),
            media_type="text/plain; charset=utf-8",
            headers={"Content-Disposition": f'attachment; filename="transcript.srt"'},
        )

    elif fmt == "json":
        return JSONResponse({
            "video_id":  video_id,
            "title":     video.title if video else "",
            "full_text": transcript.full_text,
            "segments":  segments,
            "language":  transcript.language_detected,
            "summary":   json.loads(transcript.summary) if transcript.summary else None,
            "chapters":  json.loads(transcript.chapters_json) if transcript.chapters_json else None,
        })

    elif fmt == "docx":
        return _export_docx(video, transcript, segments)

    else:
        raise HTTPException(400, "صيغة غير مدعومة. الخيارات: txt, srt, json, docx")


def _export_docx(video, transcript, segments):
    """تصدير كملف Word مع تنسيق احترافي"""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        import io

        doc = Document()

        title = doc.add_heading(video.title if video else "نص مفرَّغ", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        info = doc.add_paragraph()
        info.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        info.add_run(f"التاريخ: {video.created_at.strftime('%Y/%m/%d') if video else ''} | ")
        info.add_run(f"اللغة: {transcript.language_detected or 'عربي'}")

        doc.add_paragraph("─" * 50)

        if transcript.summary:
            try:
                summary_data = json.loads(transcript.summary)
                doc.add_heading("الملخص", 1)
                p = doc.add_paragraph(summary_data.get("summary", ""))
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

                if summary_data.get("key_points"):
                    doc.add_heading("النقاط الرئيسية", 2)
                    for point in summary_data["key_points"]:
                        p = doc.add_paragraph(f"• {point}", style="List Bullet")
                        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

                doc.add_paragraph("─" * 50)
            except Exception:
                pass

        doc.add_heading("النص المفرَّغ الكامل", 1)

        current_speaker = None
        for seg in segments:
            speaker = seg.get("speaker")
            time_str = f"[{_fmt_time(seg['start'])}]"

            if speaker and speaker != current_speaker:
                current_speaker = speaker
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                run = p.add_run(f"\n{speaker}:")
                run.bold = True
                run.font.color.rgb = RGBColor(0x34, 0xD3, 0x99)

            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            p.add_run(f"{time_str} ").font.color.rgb = RGBColor(0x9C, 0xA3, 0xAF)
            p.add_run(seg["text"])

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)

        return StreamingResponse(
            buf,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f'attachment; filename="transcript.docx"'},
        )

    except ImportError:
        raise HTTPException(400, "python-docx غير مثبت. شغّل: pip install python-docx")


def _srt_time(s: float) -> str:
    h  = int(s // 3600)
    m  = int((s % 3600) // 60)
    sc = int(s % 60)
    ms = int((s - int(s)) * 1000)
    return f"{h:02d}:{m:02d}:{sc:02d},{ms:03d}"

def _fmt_time(s: float) -> str:
    m = int(s // 60)
    return f"{m:02d}:{int(s % 60):02d}"
